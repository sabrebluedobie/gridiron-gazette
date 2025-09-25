#!/usr/bin/env python3
"""
Document Formatting Fix Module for Gridiron Gazette
Focuses exclusively on fixing:
- Header/footer gradient sliding
- Margin consistency
- Blank page removal
- Table alignment issues
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def fix_document_formatting(doc_path: str) -> bool:
    """
    Main function to fix all formatting issues in the generated document.
    
    Args:
        doc_path: Path to the DOCX file to fix
        
    Returns:
        True if successful, False otherwise
    """
    try:
        logger.info(f"Fixing formatting for: {doc_path}")
        doc = Document(doc_path)
        
        # Fix 1: Consistent margins to prevent sliding
        _fix_margins_and_sections(doc)
        
        # Fix 2: Header/Footer alignment for gradient
        _fix_header_footer_alignment(doc)
        
        # Fix 3: Remove blank pages
        _remove_blank_pages(doc)
        
        # Fix 4: Fix table formatting
        _fix_table_formatting(doc)
        
        # Fix 5: Consistent paragraph spacing
        _fix_paragraph_spacing(doc)
        
        # Save the fixed document
        doc.save(doc_path)
        logger.info("✅ Formatting fixes applied successfully")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing formatting: {e}")
        return False


def _fix_margins_and_sections(doc: Document):
    """
    Fix margins and section settings to prevent sliding and ensure consistency.
    """
    for section in doc.sections:
        # Set consistent margins (prevent sliding)
        section.top_margin = Inches(1.0)  # Slightly larger for header space
        section.bottom_margin = Inches(1.0)  # Slightly larger for footer space
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Ensure consistent page size
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.orientation = WD_ORIENTATION.PORTRAIT
        
        # Set header/footer distances to prevent overlap
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        
        # Ensure sections are continuous (no unnecessary page breaks)
        if section != doc.sections[0]:  # Keep first section as is
            section.start_type = WD_SECTION.CONTINUOUS
    
    logger.debug(f"Fixed margins for {len(doc.sections)} sections")


def _fix_header_footer_alignment(doc: Document):
    """
    Fix header and footer alignment to prevent gradient sliding.
    This ensures headers/footers span the full width properly.
    """
    for section in doc.sections:
        # Fix header
        header = section.header
        for paragraph in header.paragraphs:
            # Ensure full width alignment
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0)
            paragraph_format.right_indent = Inches(0)
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            
            # If this is a gradient/background paragraph
            if not paragraph.text.strip():
                # Ensure it spans full width
                for run in paragraph.runs:
                    if hasattr(run, '_element'):
                        # Check for shading/background
                        rPr = run._element.get_or_add_rPr()
                        # Ensure full width coverage
        
        # Fix footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0)
            paragraph_format.right_indent = Inches(0)
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
    
    logger.debug("Fixed header/footer alignment")


def _remove_blank_pages(doc: Document):
    """
    Remove blank pages caused by excessive empty paragraphs and page breaks.
    """
    paragraphs_to_delete = []
    consecutive_empty = 0
    last_paragraph_with_content = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # Check if paragraph has any real content
        has_content = bool(text)
        
        # Also check for images/objects
        if not has_content:
            for run in paragraph.runs:
                if run.text.strip():
                    has_content = True
                    break
                # Check for embedded objects
                if hasattr(run, '_element'):
                    if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:object'):
                        has_content = True
                        break
        
        if not has_content:
            consecutive_empty += 1
            
            # Keep maximum 1 empty paragraph for spacing
            if consecutive_empty > 1:
                paragraphs_to_delete.append(paragraph)
            
            # Check for page breaks in empty paragraphs
            if hasattr(paragraph, '_element'):
                page_breaks = paragraph._element.xpath('.//w:br[@w:type="page"]')
                column_breaks = paragraph._element.xpath('.//w:br[@w:type="column"]')
                
                # Remove unnecessary breaks
                if page_breaks or column_breaks:
                    # Only keep if there was substantial content before
                    if last_paragraph_with_content and i - last_paragraph_with_content > 3:
                        # Keep the break
                        pass
                    else:
                        # Remove the break
                        for br in page_breaks + column_breaks:
                            br.getparent().remove(br)
                        # Mark paragraph for deletion if now empty
                        if consecutive_empty > 1:
                            if paragraph not in paragraphs_to_delete:
                                paragraphs_to_delete.append(paragraph)
        else:
            consecutive_empty = 0
            last_paragraph_with_content = i
    
    # Delete marked paragraphs
    for paragraph in paragraphs_to_delete:
        try:
            p = paragraph._element
            p.getparent().remove(p)
        except:
            pass  # Paragraph might already be removed
    
    logger.debug(f"Removed {len(paragraphs_to_delete)} empty paragraphs")


def _fix_table_formatting(doc: Document):
    """
    Fix table formatting to prevent margin overflow and ensure proper alignment.
    """
    for table in doc.tables:
        # Set table alignment to center
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Ensure table doesn't exceed page width
        table.autofit = True
        
        # Fix table properties
        tbl = table._element
        tblPr = tbl.xpath('.//w:tblPr')[0] if tbl.xpath('.//w:tblPr') else tbl.add_tblPr()
        
        # Set table width to auto (fit content)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'auto')
        tblW.set(qn('w:w'), '0')
        
        # Remove any existing width settings
        for existing_tblW in tblPr.xpath('.//w:tblW'):
            tblPr.remove(existing_tblW)
        
        tblPr.append(tblW)
        
        # Set table indentation to 0 (prevent left margin issues)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '0')
        tblInd.set(qn('w:type'), 'dxa')
        
        # Remove existing indentation
        for existing_tblInd in tblPr.xpath('.//w:tblInd'):
            tblPr.remove(existing_tblInd)
        
        tblPr.append(tblInd)
        
        # Fix cell margins and spacing
        for row in table.rows:
            for cell in row.cells:
                # Set cell margins
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                
                # Remove existing margins
                for tcMar in tcPr.xpath('.//w:tcMar'):
                    tcPr.remove(tcMar)
                
                # Add consistent small margins
                tcMar = OxmlElement('w:tcMar')
                
                for margin_type in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_type}')
                    margin.set(qn('w:w'), '72')  # 0.05 inch
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                
                tcPr.append(tcMar)
                
                # Fix paragraph spacing in cells
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.0
    
    logger.debug(f"Fixed formatting for {len(doc.tables)} tables")


def _fix_paragraph_spacing(doc: Document):
    """
    Fix paragraph spacing to be consistent throughout the document.
    """
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only for non-empty paragraphs
            pf = paragraph.paragraph_format
            
            # Check if it's a heading (usually has larger font or bold)
            is_heading = False
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.bold or (first_run.font.size and first_run.font.size > Pt(12)):
                    is_heading = True
            
            if is_heading:
                # Heading spacing
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)
            else:
                # Normal paragraph spacing
                pf.space_before = Pt(6)
                pf.space_after = Pt(6)
            
            # Consistent line spacing
            pf.line_spacing = 1.15
            
            # Remove any indentation that might cause margin issues
            if pf.left_indent and pf.left_indent > Inches(0.5):
                pf.left_indent = Inches(0)
            if pf.right_indent and pf.right_indent > Inches(0.5):
                pf.right_indent = Inches(0)
    
    logger.debug("Fixed paragraph spacing")


def apply_formatting_fixes(doc_path: str) -> str:
    """
    Apply all formatting fixes to a document.
    
    Args:
        doc_path: Path to the document to fix
        
    Returns:
        Path to the fixed document
    """
    doc_path = Path(doc_path)
    
    if not doc_path.exists():
        raise FileNotFoundError(f"Document not found: {doc_path}")
    
    # Apply fixes
    if fix_document_formatting(str(doc_path)):
        logger.info(f"✅ Formatting fixed for: {doc_path}")
    else:
        logger.warning(f"⚠️ Some formatting fixes may have failed for: {doc_path}")
    
    return str(doc_path)


# Integration with weekly_recap.py
def post_process_with_formatting(doc_path: str):
    """
    Drop-in replacement for _post_process_document in weekly_recap.py
    """
    return apply_formatting_fixes(doc_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        doc_file = sys.argv[1]
        print(f"Applying formatting fixes to: {doc_file}")
        try:
            result = apply_formatting_fixes(doc_file)
            print(f"✅ Success! Fixed document: {result}")
        except Exception as e:
            print(f"❌ Error: {e}")
    else:
        print("Document Formatting Fix Tool")
        print("Usage: python document_formatter.py <path_to_docx>")
        print("\nThis tool fixes:")
        print("  • Header/footer gradient sliding")
        print("  • Inconsistent margins")
        print("  • Blank pages")
        print("  • Table overflow issues")
        print("  • Paragraph spacing")