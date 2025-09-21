# footer_gradient.py
from docx import Document
from docx.shared import Mm

def add_footer_gradient(docx_path, gradient_png, bar_height_mm: float = 12.0) -> None:
    """
    Adds a stable diagonal gradient bar in the footer by:
      - clearing legacy floating shapes/paragraphs
      - creating a 2-row, 1-col table with explicit width
      - inserting the gradient as an inline picture (row 1)
      - leaving row 2 for text (or blank)
    """
    doc = Document(str(docx_path))

    for section in doc.sections:
        # Keep footer from nudging
        section.bottom_margin = Mm(15)      # ~0.59"
        section.footer_distance = Mm(8)     # ~0.31"
        section.different_first_page_header_footer = False

        footer = section.footer

        # Remove any legacy footer paragraphs/shapes
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        # Compute content width (page width − left/right margins)
        content_width = section.page_width - section.left_margin - section.right_margin

        # IMPORTANT: pass width explicitly (required on some python-docx versions)
        tbl = footer.add_table(rows=2, cols=1, width=content_width)
        tbl.autofit = False

        # Force the single column to fill the content width
        tbl.columns[0].width = content_width
        for row in tbl.rows:
            row.cells[0].width = content_width

        # Row 1: gradient strip (inline, stable)
        cell = tbl.rows[0].cells[0]
        run = cell.paragraphs[0].add_run()
        try:
            pic = run.add_picture(str(gradient_png))
            pic.height = Mm(bar_height_mm)   # width auto-fits to the cell
        except Exception as e:
            print(f"[footer] Could not add gradient image: {e}")

        # Row 2: reserved for footer text (or leave blank if template provides it)
        tbl.rows[1].cells[0].paragraphs[0].add_run("")

    doc.save(str(docx_path))
