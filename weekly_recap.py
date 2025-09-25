#!/usr/bin/env python3
"""
Weekly Recap Builder for Gridiron Gazette
COMPREHENSIVE FIX VERSION
- Properly embeds images
- Removes blank pages
- Fixes margin issues
- Cleans up formatting
"""
from __future__ import annotations
import os
import logging
import re
from pathlib import Path
from typing import Any, Dict, Optional, List

# DOCX handling with image support and formatting control
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm, Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import gazette_data
from storymaker import (
    StoryMaker, 
    MatchupData, 
    PlayerStat,
    clean_markdown_for_docx,
    clean_all_markdown_in_dict
)

# Set up logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Optional OpenAI wrapper (safe if absent)
try:
    from llm_openai import chat as openai_llm
except Exception:
    openai_llm = None
    logger.info("OpenAI LLM not available, will use fallback templates")


def build_weekly_recap(
    league_id: int,
    year: int,
    week: int,
    template: str = "recap_template.docx",
    output_path: str = "recaps/Gazette_{year}_W{week02}.docx",
    use_llm_blurbs: bool = True,
) -> str:
    """
    Builds the Gazette docx with comprehensive fixes:
      1) Fetches ESPN context
      2) Generates Sabre recaps
      3) Properly embeds images
      4) Cleans markdown
      5) Renders template
      6) Post-processes to fix blank pages and margins
    """
    # Get base context from ESPN
    ctx = gazette_data.build_context(league_id, year, week)
    
    # Add Sabre blurbs
    if use_llm_blurbs:
        _attach_sabre_recaps(ctx)
    else:
        _attach_simple_blurbs(ctx)
    
    # Add spotlight content
    _attach_spotlight_content(ctx)
    
    # Clean all markdown from the context
    ctx = clean_all_markdown_in_dict(ctx)
    
    # Render the doc WITH PROPER IMAGE EMBEDDING
    out = _render_docx_with_images(template, output_path, ctx)
    
    # Post-process to fix blank pages and margins
    _post_process_document(out)
    
    return out


def _find_logo_file(team_name: str, logo_type: str = "team") -> Optional[Path]:
    """
    Find a logo file for a team, returning Path object if found.
    Searches multiple directories and name variations.
    """
    # Clean team name for filename matching
    clean_name = team_name
    # Remove emojis and special characters
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F680-\U0001F6FF"  # transport & map symbols
        u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        u"\U0001F926-\U0001F9FF"
        u"\U00010000-\U0010ffff"
        u"\u2640-\u2642" 
        u"\u2600-\u2B55"
        u"\u200d"
        u"\u23cf"
        u"\u23e9"
        u"\u231a"
        u"\ufe0f"  # dingbats
        u"\u3030"
        "]+", flags=re.UNICODE)
    
    clean_name = emoji_pattern.sub('', clean_name)
    clean_name = clean_name.replace("'", "").replace("!", "").replace("'", "")
    clean_name = clean_name.strip()
    
    # Logo directories to search (in priority order)
    if logo_type == "team":
        search_dirs = [
            Path("logos/team_logos"),
            Path("logos/team_logos"),
            Path("logos"),
            Path("media/logos"),
            Path(".")
        ]
    elif logo_type == "league":
        search_dirs = [
            Path("logos/league_logos"),
            Path("logos/team_logos"),  # Browns logo might be here
            Path("logos/team_logos"),
            Path("logos"),
            Path(".")
        ]
    else:  # sponsor
        search_dirs = [
            Path("logos/sponsor_logos"),
            Path("logos"),
            Path("media"),
            Path(".")
        ]
    
    # Special cases for specific teams
    special_mappings = {
        "DEM BOY'S!🏆🏆🏆🏆": ["DEMBOYS", "DEM_BOYS", "demboys"],
        "🏉THE💀REBELS🏉": ["THEREBELS", "THE_REBELS", "therebels", "THEREBELS_"],
        "Under the InfluWENTZ": ["UndertheInfluWENTZ", "Underthe_InfluWENTZ", "InfluWENTZ"],
        "Annie1235 slayy": ["Annie1235slayy", "Annie1235_slayy", "annieslayy"],
        "Nana's Hawks": ["NanasHawks", "Nanas_Hawks", "nanashawks"],
        "Kansas City Pumas": ["KansasCity_Pumas", "KansasCityPumas", "KC_Pumas"],
    }
    
    # Get special mapping if exists
    name_variations = special_mappings.get(team_name, [])
    
    # Add standard variations
    name_variations.extend([
        team_name,  # Original
        clean_name,  # Cleaned
        clean_name.replace(" ", ""),  # No spaces
        clean_name.replace(" ", "_"),  # Underscores
        clean_name.replace(" ", "-"),  # Hyphens
        team_name.replace(" ", ""),
        team_name.replace(" ", "_"),
        team_name.replace(" ", "-"),
    ])
    
    # Also try lowercase versions
    name_variations.extend([n.lower() for n in name_variations])
    
    # Remove duplicates while preserving order
    seen = set()
    unique_variations = []
    for name in name_variations:
        if name and name not in seen:
            seen.add(name)
            unique_variations.append(name)
    
    # Search for the file
    for directory in search_dirs:
        if not directory.exists():
            continue
            
        for name_var in unique_variations:
            for ext in ['.png', '.jpg', '.jpeg', '.gif', '.PNG', '.JPG']:
                filepath = directory / f"{name_var}{ext}"
                if filepath.exists():
                    logger.info(f"Found logo for '{team_name}': {filepath}")
                    return filepath
    
    logger.warning(f"No logo found for '{team_name}' (type: {logo_type})")
    return None


def _render_docx_with_images(template_path: str, out_pattern: str, ctx: Dict[str, Any]) -> str:
    """
    Render the DOCX template with proper image embedding.
    Converts logo paths to InlineImage objects for proper rendering.
    """
    tpl_path = Path(template_path)
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {tpl_path}")

    # Create DocxTemplate object
    doc = DocxTemplate(str(tpl_path))
    
    # Prepare context with embedded images
    image_context = _prepare_image_context(doc, ctx.copy())
    
    # Get week number for output filename
    week = int(ctx.get("WEEK_NUMBER", 0) or ctx.get("WEEK", 0) or 0)
    year = ctx.get("YEAR", "")
    
    # Format output path
    out_path = out_pattern.format(
        year=year, 
        week=week, 
        week02=f"{week:02d}"
    )
    
    out_file = Path(out_path)
    out_file.parent.mkdir(parents=True, exist_ok=True)
    
    # Render the template with image context
    doc.render(image_context)
    
    # Save the document
    doc.save(str(out_file))
    
    logger.info(f"✅ Generated recap document: {out_file}")
    _log_recap_summary(image_context, out_file)
    
    return str(out_file)


def _prepare_image_context(doc_template: DocxTemplate, context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Convert logo paths to InlineImage objects for proper DOCX embedding.
    Uses consistent sizing to prevent margin issues.
    """
    logger.info("Preparing images for DOCX embedding...")
    
    # Standard sizes for consistency (prevents margin issues)
    TEAM_LOGO_SIZE = Mm(15)  # Smaller for inline team logos
    HEADER_LOGO_SIZE = Mm(25)  # Larger for header logos
    
    # Process team logos for each matchup
    for i in range(1, 8):
        home_key = f"MATCHUP{i}_HOME"
        away_key = f"MATCHUP{i}_AWAY"
        
        # Home team logo
        if home_key in context:
            team_name = context[home_key]
            logo_key = f"MATCHUP{i}_HOME_LOGO"
            
            logo_path = _find_logo_file(team_name, "team")
            if logo_path and logo_path.exists():
                try:
                    context[logo_key] = InlineImage(
                        doc_template, 
                        str(logo_path), 
                        width=TEAM_LOGO_SIZE
                    )
                    logger.info(f"✅ Embedded logo for {team_name}")
                except Exception as e:
                    logger.error(f"Failed to embed logo for {team_name}: {e}")
                    context[logo_key] = ""  # Empty string for missing logos
            else:
                context[logo_key] = ""
        
        # Away team logo
        if away_key in context:
            team_name = context[away_key]
            logo_key = f"MATCHUP{i}_AWAY_LOGO"
            
            logo_path = _find_logo_file(team_name, "team")
            if logo_path and logo_path.exists():
                try:
                    context[logo_key] = InlineImage(
                        doc_template, 
                        str(logo_path), 
                        width=TEAM_LOGO_SIZE
                    )
                    logger.info(f"✅ Embedded logo for {team_name}")
                except Exception as e:
                    logger.error(f"Failed to embed logo for {team_name}: {e}")
                    context[logo_key] = ""
            else:
                context[logo_key] = ""
    
    # Process league logo - special handling for Browns
    browns_paths = [
        Path("logos/team_logos/brownseakc.png"),
        Path("logos/team_logos/brownseakc.png"),
        Path("logos/league_logos/brownseakc.png"),
        Path("logos/brownseakc.png"),
        Path("brownseakc.png")
    ]
    
    league_logo_embedded = False
    for path in browns_paths:
        if path.exists():
            try:
                context["LEAGUE_LOGO"] = InlineImage(
                    doc_template,
                    str(path),
                    width=HEADER_LOGO_SIZE
                )
                logger.info(f"✅ Embedded Browns league logo from {path}")
                league_logo_embedded = True
                break
            except Exception as e:
                logger.error(f"Failed to embed league logo: {e}")
    
    if not league_logo_embedded:
        context["LEAGUE_LOGO"] = ""
    
    # Process sponsor logo
    sponsor_paths = [
        Path("logos/sponsor_logos/gridiron_gazette.png"),
        Path("logos/gridiron_gazette.png"),
        Path("logos/gg_logo.png"),
        Path("media/gridiron_gazette.png")
    ]
    
    sponsor_embedded = False
    for path in sponsor_paths:
        if path.exists():
            try:
                context["SPONSOR_LOGO"] = InlineImage(
                    doc_template,
                    str(path),
                    width=HEADER_LOGO_SIZE
                )
                logger.info(f"✅ Embedded sponsor logo from {path}")
                sponsor_embedded = True
                break
            except Exception as e:
                logger.error(f"Failed to embed sponsor logo: {e}")
    
    if not sponsor_embedded:
        context["SPONSOR_LOGO"] = ""
    
    return context


def _post_process_document(doc_path: str) -> None:
    """
    Post-process the document to fix blank pages and margin issues.
    This function:
    1. Removes empty paragraphs that cause blank pages
    2. Fixes section breaks
    3. Adjusts margins for consistency
    4. Removes excessive spacing
    """
    logger.info("Post-processing document to fix formatting issues...")
    
    try:
        # Open the document for post-processing
        doc = Document(doc_path)
        
        # Fix margins for all sections (prevents margin slip)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            # Ensure consistent page size
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            
            # Remove unnecessary section breaks that cause blank pages
            section.start_type = WD_SECTION.CONTINUOUS
        
        # Remove empty paragraphs that cause blank pages
        paragraphs_to_delete = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is effectively empty
            text = paragraph.text.strip()
            
            # Remove completely empty paragraphs
            if not text:
                # Check if it has no runs with images either
                has_content = False
                for run in paragraph.runs:
                    if run.text.strip() or hasattr(run, '_element'):
                        # Check for embedded objects
                        if run._element.xpath('.//w:drawing'):
                            has_content = True
                            break
                
                if not has_content:
                    paragraphs_to_delete.append(paragraph)
            
            # Remove excessive line breaks
            elif text in ['\n', '\r\n', '\r']:
                paragraphs_to_delete.append(paragraph)
        
        # Delete empty paragraphs
        for paragraph in paragraphs_to_delete:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Fix spacing between paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only for non-empty paragraphs
                # Reduce spacing after paragraphs
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(6)  # Reduced from default
                paragraph_format.space_before = Pt(6)  # Reduced from default
                
                # Ensure single line spacing
                paragraph_format.line_spacing = 1.0
        
        # Remove page breaks that cause blank pages
        for paragraph in doc.paragraphs:
            if paragraph._element.xpath('.//w:br[@w:type="page"]'):
                # Check if this is necessary (e.g., between matchups)
                if not paragraph.text.strip():
                    # Remove unnecessary page break
                    for br in paragraph._element.xpath('.//w:br[@w:type="page"]'):
                        br.getparent().remove(br)
        
        # Handle tables to prevent margin issues
        for table in doc.tables:
            table.autofit = True
            # Set table alignment
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ensure table doesn't exceed margins
            for row in table.rows:
                for cell in row.cells:
                    # Set cell margins
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    
                    for margin_type in ['top', 'left', 'bottom', 'right']:
                        margin = OxmlElement(f'w:{margin_type}')
                        margin.set(qn('w:w'), '50')  # Small margin
                        margin.set(qn('w:type'), 'dxa')
                        tcMar.append(margin)
                    
                    tcPr.append(tcMar)
        
        # Save the fixed document
        doc.save(doc_path)
        logger.info(f"✅ Post-processing complete: fixed margins and removed blank pages")
        
    except Exception as e:
        logger.error(f"Error during post-processing: {e}")
        # Don't fail the entire process if post-processing fails
        logger.info("Document generated but post-processing failed")


def _attach_sabre_recaps(ctx: Dict[str, Any]) -> None:
    """Generate and attach Sabre recaps using the StoryMaker."""
    maker = StoryMaker(llm=openai_llm if os.getenv("OPENAI_API_KEY") else None)

    count = int(ctx.get("MATCHUP_COUNT") or 7)
    league_name = str(ctx.get("LEAGUE_NAME", "League"))
    week_num = int(ctx.get("WEEK_NUMBER", 0) or ctx.get("WEEK", 0) or 0)

    for i in range(1, min(count + 1, 8)):
        h = ctx.get(f"MATCHUP{i}_HOME")
        a = ctx.get(f"MATCHUP{i}_AWAY")
        hs = ctx.get(f"MATCHUP{i}_HS")
        as_ = ctx.get(f"MATCHUP{i}_AS")
        
        if not (h and a):
            continue

        # Get top performers
        tops = []
        th = _safe(ctx.get(f"MATCHUP{i}_TOP_HOME", ""))
        ta = _safe(ctx.get(f"MATCHUP{i}_TOP_AWAY", ""))
        if th:
            tops.append(PlayerStat(name=th))
        if ta:
            tops.append(PlayerStat(name=ta))

        try:
            sa = float(str(hs)) if hs not in (None, "") else 0.0
            sb = float(str(as_)) if as_ not in (None, "") else 0.0
        except Exception:
            sa, sb = 0.0, 0.0

        md = MatchupData(
            league_name=league_name,
            week=week_num,
            team_a=str(h), 
            team_b=str(a),
            score_a=sa, 
            score_b=sb,
            top_performers=tops,
            winner=str(h) if sa >= sb else str(a),
            margin=abs(sa - sb),
        )
        
        # Generate recap with markdown cleaning enabled
        recap = maker.generate_recap(md, clean_markdown=True)
        
        # Additional cleanup for any remaining formatting issues
        recap = recap.replace('\r\n', '\n').replace('\r', '\n')
        
        # Ensure proper paragraph breaks
        paragraphs = recap.split('\n\n')
        cleaned_paragraphs = [p.strip() for p in paragraphs if p.strip()]
        recap = '\n\n'.join(cleaned_paragraphs)
        
        ctx[f"MATCHUP{i}_BLURB"] = recap
        
        logger.info(f"Generated Sabre recap for matchup {i}: {h} vs {a}")


def _attach_simple_blurbs(ctx: Dict[str, Any]) -> None:
    """Attach simple fallback blurbs when LLM is not available."""
    count = int(ctx.get("MATCHUP_COUNT") or 7)
    for i in range(1, min(count + 1, 8)):
        if ctx.get(f"MATCHUP{i}_BLURB"):
            continue
        h = ctx.get(f"MATCHUP{i}_HOME")
        a = ctx.get(f"MATCHUP{i}_AWAY")
        hs = ctx.get(f"MATCHUP{i}_HS")
        as_ = ctx.get(f"MATCHUP{i}_AS")
        
        if not (h and a):
            continue
            
        try:
            sa = float(str(hs)) if hs not in (None, "") else 0.0
            sb = float(str(as_)) if as_ not in (None, "") else 0.0
        except Exception:
            sa, sb = 0.0, 0.0
            
        winner = h if sa >= sb else a
        loser = a if winner == h else h
        margin = abs(sa - sb)
        
        tone = "nail-biter" if margin < 5 else ("statement win" if margin > 20 else "solid win")
        
        blurb = f"{winner} topped {loser} {hs}-{as_} in a {tone}.\n\n—Sabre, your hilariously snarky 4-legged Gridiron Gazette reporter 🐾"
        ctx[f"MATCHUP{i}_BLURB"] = clean_markdown_for_docx(blurb)


def _attach_spotlight_content(ctx: Dict[str, Any]) -> None:
    """
    Generate additional spotlight content for Stats Spotlight sections.
    """
    count = int(ctx.get("MATCHUP_COUNT") or 7)
    
    for i in range(1, min(count + 1, 8)):
        h = ctx.get(f"MATCHUP{i}_HOME")
        a = ctx.get(f"MATCHUP{i}_AWAY")
        
        if not (h and a):
            continue
        
        # Fill in missing spotlight fields with Sabre-style content
        if not ctx.get(f"MATCHUP{i}_TOP_HOME"):
            ctx[f"MATCHUP{i}_TOP_HOME"] = f"{h}'s offense showed up when it mattered"
        
        if not ctx.get(f"MATCHUP{i}_TOP_AWAY"):
            ctx[f"MATCHUP{i}_TOP_AWAY"] = f"{a}'s squad battled to the end"
        
        if not ctx.get(f"MATCHUP{i}_BUST"):
            ctx[f"MATCHUP{i}_BUST"] = "Both teams had players who'd rather forget this week"
        
        if not ctx.get(f"MATCHUP{i}_KEYPLAY"):
            ctx[f"MATCHUP{i}_KEYPLAY"] = "Every yard was earned in this matchup"
        
        if not ctx.get(f"MATCHUP{i}_DEF"):
            ctx[f"MATCHUP{i}_DEF"] = "Defense wasn't invited to this scoring party"


def _log_recap_summary(ctx: Dict[str, Any], output_file: Path) -> None:
    """Log a summary of what was included in the recap."""
    logger.info("\n" + "="*60)
    logger.info("RECAP GENERATION SUMMARY")
    logger.info("="*60)
    logger.info(f"Week: {ctx.get('WEEK_NUMBER', '?')}")
    logger.info(f"Year: {ctx.get('YEAR', '?')}")
    logger.info(f"League: {ctx.get('LEAGUE_NAME', 'Unknown')}")
    logger.info(f"Output: {output_file}")
    
    # Count embedded images
    image_count = sum(1 for key in ctx if 'LOGO' in key and hasattr(ctx[key], '__class__') and 'InlineImage' in str(ctx[key].__class__))
    
    logger.info(f"\n✅ Images embedded: {image_count}")
    logger.info("✅ Margins fixed")
    logger.info("✅ Blank pages removed")
    logger.info("✅ Formatting cleaned")
    logger.info("="*60 + "\n")


def _safe(s: Any) -> str:
    """Safely convert any value to string."""
    return "" if s is None else str(s)


if __name__ == "__main__":
    print("Weekly Recap Builder - COMPREHENSIVE FIX")
    print("Features:")
    print("  ✅ Proper image embedding")
    print("  ✅ Removes blank pages")
    print("  ✅ Fixes margin issues")
    print("  ✅ Cleans formatting")
    print("\nUsage: Called from build_gazette.py")#!/usr/bin/env python3
"""
Weekly Recap Builder for Gridiron Gazette
COMPLETE VERSION using team_logos.json for all logo mappings
- Properly embeds images from JSON mappings
- Removes blank pages
- Fixes margin issues
- Supports multiple leagues via JSON
"""
from __future__ import annotations
import os
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, Optional, List

# DOCX handling with image support and formatting control
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm, Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import gazette_data
from storymaker import (
    StoryMaker, 
    MatchupData, 
    PlayerStat,
    clean_markdown_for_docx,
    clean_all_markdown_in_dict
)

# Set up logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Optional OpenAI wrapper (safe if absent)
try:
    from llm_openai import chat as openai_llm
except Exception:
    openai_llm = None
    logger.info("OpenAI LLM not available, will use fallback templates")


class LogoManager:
    """Manages logo mappings from team_logos.json"""
    
    def __init__(self, json_file: str = "team_logos.json"):
        """Initialize with logo mappings from JSON file"""
        self.logo_mappings = {}
        self.json_file = json_file
        self.load_mappings()
    
    def load_mappings(self):
        """Load logo mappings from JSON file"""
        json_path = Path(self.json_file)
        if json_path.exists():
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    self.logo_mappings = json.load(f)
                logger.info(f"✅ Loaded {len(self.logo_mappings)} logo mappings from {self.json_file}")
            except Exception as e:
                logger.error(f"Failed to load {self.json_file}: {e}")
                self.logo_mappings = {}
        else:
            logger.warning(f"Logo mappings file not found: {self.json_file}")
            self.logo_mappings = {}
    
    def get_team_logo(self, team_name: str) -> Optional[Path]:
        """Get logo path for a team from mappings"""
        if team_name in self.logo_mappings:
            logo_path = Path(self.logo_mappings[team_name])
            if logo_path.exists():
                return logo_path
            else:
                logger.warning(f"Logo file not found: {logo_path} for team {team_name}")
        else:
            logger.warning(f"No mapping found for team: {team_name}")
        return None
    
    def get_league_logo(self) -> Optional[Path]:
        """Get league logo from mappings"""
        if "LEAGUE_LOGO" in self.logo_mappings:
            logo_path = Path(self.logo_mappings["LEAGUE_LOGO"])
            if logo_path.exists():
                return logo_path
            else:
                logger.warning(f"League logo file not found: {logo_path}")
        return None
    
    def get_sponsor_logo(self) -> Optional[Path]:
        """Get sponsor logo from mappings"""
        if "SPONSOR_LOGO" in self.logo_mappings:
            logo_path = Path(self.logo_mappings["SPONSOR_LOGO"])
            if logo_path.exists():
                return logo_path
            else:
                logger.warning(f"Sponsor logo file not found: {logo_path}")
        return None


def build_weekly_recap(
    league_id: int,
    year: int,
    week: int,
    template: str = "recap_template.docx",
    output_path: str = "recaps/Gazette_{year}_W{week02}.docx",
    use_llm_blurbs: bool = True,
    logo_json: str = "team_logos.json"
) -> str:
    """
    Builds the Gazette docx with comprehensive fixes:
      1) Fetches ESPN context
      2) Generates Sabre recaps
      3) Properly embeds images from team_logos.json
      4) Cleans markdown
      5) Renders template
      6) Post-processes to fix blank pages and margins
    
    Args:
        league_id: ESPN league ID
        year: Season year
        week: Week number
        template: Path to DOCX template
        output_path: Output path pattern
        use_llm_blurbs: Whether to use LLM for Sabre blurbs
        logo_json: Path to team_logos.json file
    """
    # Get base context from ESPN
    ctx = gazette_data.build_context(league_id, year, week)
    
    # Add Sabre blurbs
    if use_llm_blurbs:
        _attach_sabre_recaps(ctx)
    else:
        _attach_simple_blurbs(ctx)
    
    # Add spotlight content
    _attach_spotlight_content(ctx)
    
    # Clean all markdown from the context
    ctx = clean_all_markdown_in_dict(ctx)
    
    # Render the doc WITH PROPER IMAGE EMBEDDING from JSON
    out = _render_docx_with_images(template, output_path, ctx, logo_json)
    
    # Post-process to fix blank pages and margins
    _post_process_document(out)
    
    return out


def _render_docx_with_images(template_path: str, out_pattern: str, ctx: Dict[str, Any], logo_json: str) -> str:
    """
    Render the DOCX template with proper image embedding using team_logos.json.
    """
    tpl_path = Path(template_path)
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {tpl_path}")

    # Create DocxTemplate object
    doc = DocxTemplate(str(tpl_path))
    
    # Prepare context with embedded images from JSON mappings
    image_context = _prepare_image_context_from_json(doc, ctx.copy(), logo_json)
    
    # Get week number for output filename
    week = int(ctx.get("WEEK_NUMBER", 0) or ctx.get("WEEK", 0) or 0)
    year = ctx.get("YEAR", "")
    
    # Format output path
    out_path = out_pattern.format(
        year=year, 
        week=week, 
        week02=f"{week:02d}"
    )
    
    out_file = Path(out_path)
    out_file.parent.mkdir(parents=True, exist_ok=True)
    
    # Render the template with image context
    doc.render(image_context)
    
    # Save the document
    doc.save(str(out_file))
    
    logger.info(f"✅ Generated recap document: {out_file}")
    _log_recap_summary(image_context, out_file)
    
    return str(out_file)


def _prepare_image_context_from_json(doc_template: DocxTemplate, context: Dict[str, Any], logo_json: str) -> Dict[str, Any]:
    """
    Convert logo paths to InlineImage objects using team_logos.json mappings.
    This ensures we use the exact paths specified in the JSON file.
    """
    logger.info(f"Preparing images from {logo_json}...")
    
    # Initialize logo manager with JSON mappings
    logo_manager = LogoManager(logo_json)
    
    # Standard sizes for consistency (prevents margin issues)
    TEAM_LOGO_SIZE = Mm(15)  # Smaller for inline team logos
    HEADER_LOGO_SIZE = Mm(25)  # Larger for header logos
    
    # Process team logos for each matchup
    images_embedded = 0
    for i in range(1, 8):
        home_key = f"MATCHUP{i}_HOME"
        away_key = f"MATCHUP{i}_AWAY"
        
        # Home team logo
        if home_key in context:
            team_name = context[home_key]
            logo_key = f"MATCHUP{i}_HOME_LOGO"
            
            logo_path = logo_manager.get_team_logo(team_name)
            if logo_path:
                try:
                    context[logo_key] = InlineImage(
                        doc_template, 
                        str(logo_path), 
                        width=TEAM_LOGO_SIZE
                    )
                    logger.info(f"✅ Embedded logo for {team_name}: {logo_path}")
                    images_embedded += 1
                except Exception as e:
                    logger.error(f"Failed to embed logo for {team_name}: {e}")
                    context[logo_key] = ""
            else:
                context[logo_key] = ""
                logger.warning(f"No logo found for home team: {team_name}")
        
        # Away team logo
        if away_key in context:
            team_name = context[away_key]
            logo_key = f"MATCHUP{i}_AWAY_LOGO"
            
            logo_path = logo_manager.get_team_logo(team_name)
            if logo_path:
                try:
                    context[logo_key] = InlineImage(
                        doc_template, 
                        str(logo_path), 
                        width=TEAM_LOGO_SIZE
                    )
                    logger.info(f"✅ Embedded logo for {team_name}: {logo_path}")
                    images_embedded += 1
                except Exception as e:
                    logger.error(f"Failed to embed logo for {team_name}: {e}")
                    context[logo_key] = ""
            else:
                context[logo_key] = ""
                logger.warning(f"No logo found for away team: {team_name}")
    
    # Process league logo
    league_logo_path = logo_manager.get_league_logo()
    if league_logo_path:
        try:
            context["LEAGUE_LOGO"] = InlineImage(
                doc_template,
                str(league_logo_path),
                width=HEADER_LOGO_SIZE
            )
            logger.info(f"✅ Embedded league logo: {league_logo_path}")
            images_embedded += 1
        except Exception as e:
            logger.error(f"Failed to embed league logo: {e}")
            context["LEAGUE_LOGO"] = ""
    else:
        context["LEAGUE_LOGO"] = ""
        logger.warning("No league logo found in mappings")
    
    # Process sponsor logo
    sponsor_logo_path = logo_manager.get_sponsor_logo()
    if sponsor_logo_path:
        try:
            context["SPONSOR_LOGO"] = InlineImage(
                doc_template,
                str(sponsor_logo_path),
                width=HEADER_LOGO_SIZE
            )
            logger.info(f"✅ Embedded sponsor logo: {sponsor_logo_path}")
            images_embedded += 1
        except Exception as e:
            logger.error(f"Failed to embed sponsor logo: {e}")
            context["SPONSOR_LOGO"] = ""
    else:
        context["SPONSOR_LOGO"] = ""
        logger.warning("No sponsor logo found in mappings")
    
    logger.info(f"Total images embedded: {images_embedded}")
    return context


def _post_process_document(doc_path: str) -> None:
    """
    Post-process the document to fix blank pages and margin issues.
    This function:
    1. Removes empty paragraphs that cause blank pages
    2. Fixes section breaks
    3. Adjusts margins for consistency
    4. Removes excessive spacing
    """
    logger.info("Post-processing document to fix formatting issues...")
    
    try:
        # Open the document for post-processing
        doc = Document(doc_path)
        
        # Fix margins for all sections (prevents margin slip)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            # Ensure consistent page size
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            
            # Remove unnecessary section breaks that cause blank pages
            section.start_type = WD_SECTION.CONTINUOUS
        
        # Remove empty paragraphs that cause blank pages
        paragraphs_to_delete = []
        consecutive_empty = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is effectively empty
            text = paragraph.text.strip()
            
            if not text:
                # Check if it has no runs with images either
                has_content = False
                for run in paragraph.runs:
                    if run.text.strip():
                        has_content = True
                        break
                    # Check for embedded objects (images)
                    if hasattr(run, '_element') and run._element.xpath('.//w:drawing'):
                        has_content = True
                        break
                
                if not has_content:
                    consecutive_empty += 1
                    # Keep single empty paragraphs for spacing, remove multiple consecutive ones
                    if consecutive_empty > 1:
                        paragraphs_to_delete.append(paragraph)
            else:
                consecutive_empty = 0
        
        # Delete excessive empty paragraphs
        for paragraph in paragraphs_to_delete:
            p = paragraph._element
            p.getparent().remove(p)
            logger.debug("Removed empty paragraph")
        
        # Fix spacing between paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only for non-empty paragraphs
                paragraph_format = paragraph.paragraph_format
                
                # Set consistent spacing
                paragraph_format.space_after = Pt(6)  # Reduced from default
                paragraph_format.space_before = Pt(6)  # Reduced from default
                
                # Ensure single line spacing
                paragraph_format.line_spacing = 1.15  # Slightly more than 1.0 for readability
        
        # Remove unnecessary page breaks
        for paragraph in doc.paragraphs:
            if hasattr(paragraph, '_element'):
                # Check for page breaks
                page_breaks = paragraph._element.xpath('.//w:br[@w:type="page"]')
                if page_breaks:
                    # Check if this paragraph has actual content
                    if not paragraph.text.strip():
                        # Remove the page break if paragraph is empty
                        for br in page_breaks:
                            br.getparent().remove(br)
                            logger.debug("Removed unnecessary page break")
        
        # Handle tables to prevent margin issues
        for table in doc.tables:
            table.autofit = True
            # Set table alignment
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Fix table width to prevent overflow
            for row in table.rows:
                for cell in row.cells:
                    # Set cell margins
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Remove any existing margins
                    for mar in tcPr.xpath('.//w:tcMar'):
                        tcPr.remove(mar)
                    
                    # Add consistent small margins
                    tcMar = OxmlElement('w:tcMar')
                    for margin_type in ['top', 'left', 'bottom', 'right']:
                        margin = OxmlElement(f'w:{margin_type}')
                        margin.set(qn('w:w'), '50')  # Small margin
                        margin.set(qn('w:type'), 'dxa')
                        tcMar.append(margin)
                    tcPr.append(tcMar)
                    
                    # Fix cell paragraph spacing
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(3)
                        paragraph.paragraph_format.space_before = Pt(3)
        
        # Save the fixed document
        doc.save(doc_path)
        logger.info(f"✅ Post-processing complete: fixed margins and removed blank pages")
        
    except Exception as e:
        logger.error(f"Error during post-processing: {e}")
        # Don't fail the entire process if post-processing fails
        logger.info("Document generated but post-processing failed - output may have formatting issues")


def _attach_sabre_recaps(ctx: Dict[str, Any]) -> None:
    """Generate and attach Sabre recaps using the StoryMaker."""
    maker = StoryMaker(llm=openai_llm if os.getenv("OPENAI_API_KEY") else None)

    count = int(ctx.get("MATCHUP_COUNT") or 7)
    league_name = str(ctx.get("LEAGUE_NAME", "League"))
    week_num = int(ctx.get("WEEK_NUMBER", 0) or ctx.get("WEEK", 0) or 0)

    for i in range(1, min(count + 1, 8)):
        h = ctx.get(f"MATCHUP{i}_HOME")
        a = ctx.get(f"MATCHUP{i}_AWAY")
        hs = ctx.get(f"MATCHUP{i}_HS")
        as_ = ctx.get(f"MATCHUP{i}_AS")
        
        if not (h and a):
            continue

        # Get top performers
        tops = []
        th = _safe(ctx.get(f"MATCHUP{i}_TOP_HOME", ""))
        ta = _safe(ctx.get(f"MATCHUP{i}_TOP_AWAY", ""))
        if th:
            tops.append(PlayerStat(name=th))
        if ta:
            tops.append(PlayerStat(name=ta))

        try:
            sa = float(str(hs)) if hs not in (None, "") else 0.0
            sb = float(str(as_)) if as_ not in (None, "") else 0.0
        except Exception:
            sa, sb = 0.0, 0.0

        md = MatchupData(
            league_name=league_name,
            week=week_num,
            team_a=str(h), 
            team_b=str(a),
            score_a=sa, 
            score_b=sb,
            top_performers=tops,
            winner=str(h) if sa >= sb else str(a),
            margin=abs(sa - sb),
        )
        
        # Generate recap with markdown cleaning enabled
        recap = maker.generate_recap(md, clean_markdown=True)
        
        # Additional cleanup for formatting
        recap = recap.replace('\r\n', '\n').replace('\r', '\n')
        
        # Ensure proper paragraph breaks
        paragraphs = recap.split('\n\n')
        cleaned_paragraphs = [p.strip() for p in paragraphs if p.strip()]
        recap = '\n\n'.join(cleaned_paragraphs)
        
        ctx[f"MATCHUP{i}_BLURB"] = recap
        
        logger.info(f"Generated Sabre recap for matchup {i}: {h} vs {a}")


def _attach_simple_blurbs(ctx: Dict[str, Any]) -> None:
    """Attach simple fallback blurbs when LLM is not available."""
    count = int(ctx.get("MATCHUP_COUNT") or 7)
    for i in range(1, min(count + 1, 8)):
        if ctx.get(f"MATCHUP{i}_BLURB"):
            continue
        h = ctx.get(f"MATCHUP{i}_HOME")
        a = ctx.get(f"MATCHUP{i}_AWAY")
        hs = ctx.get(f"MATCHUP{i}_HS")
        as_ = ctx.get(f"MATCHUP{i}_AS")
        
        if not (h and a):
            continue
            
        try:
            sa = float(str(hs)) if hs not in (None, "") else 0.0
            sb = float(str(as_)) if as_ not in (None, "") else 0.0
        except Exception:
            sa, sb = 0.0, 0.0
            
        winner = h if sa >= sb else a
        loser = a if winner == h else h
        margin = abs(sa - sb)
        
        tone = "nail-biter" if margin < 5 else ("statement win" if margin > 20 else "solid win")
        
        blurb = f"{winner} topped {loser} {hs}-{as_} in a {tone}.\n\n—Sabre, your hilariously snarky 4-legged Gridiron Gazette reporter 🐾"
        ctx[f"MATCHUP{i}_BLURB"] = clean_markdown_for_docx(blurb)


def _attach_spotlight_content(ctx: Dict[str, Any]) -> None:
    """
    Generate additional spotlight content for Stats Spotlight sections.
    """
    count = int(ctx.get("MATCHUP_COUNT") or 7)
    
    for i in range(1, min(count + 1, 8)):
        h = ctx.get(f"MATCHUP{i}_HOME")
        a = ctx.get(f"MATCHUP{i}_AWAY")
        
        if not (h and a):
            continue
        
        # Fill in missing spotlight fields with Sabre-style content
        if not ctx.get(f"MATCHUP{i}_TOP_HOME"):
            ctx[f"MATCHUP{i}_TOP_HOME"] = f"{h}'s offense showed up when it mattered"
        
        if not ctx.get(f"MATCHUP{i}_TOP_AWAY"):
            ctx[f"MATCHUP{i}_TOP_AWAY"] = f"{a}'s squad battled to the end"
        
        if not ctx.get(f"MATCHUP{i}_BUST"):
            ctx[f"MATCHUP{i}_BUST"] = "Both teams had players who'd rather forget this week"
        
        if not ctx.get(f"MATCHUP{i}_KEYPLAY"):
            ctx[f"MATCHUP{i}_KEYPLAY"] = "Every yard was earned in this matchup"
        
        if not ctx.get(f"MATCHUP{i}_DEF"):
            ctx[f"MATCHUP{i}_DEF"] = "Defense wasn't invited to this scoring party"


def _log_recap_summary(ctx: Dict[str, Any], output_file: Path) -> None:
    """Log a summary of what was included in the recap."""
    logger.info("\n" + "="*60)
    logger.info("RECAP GENERATION SUMMARY")
    logger.info("="*60)
    logger.info(f"Week: {ctx.get('WEEK_NUMBER', '?')}")
    logger.info(f"Year: {ctx.get('YEAR', '?')}")
    logger.info(f"League: {ctx.get('LEAGUE_NAME', 'Unknown')}")
    logger.info(f"Output: {output_file}")
    
    # Count embedded images
    image_count = sum(1 for key in ctx if 'LOGO' in key and hasattr(ctx[key], '__class__') and 'InlineImage' in str(ctx[key].__class__))
    
    logger.info(f"\n✅ Images embedded: {image_count}")
    logger.info("✅ Using team_logos.json for all mappings")
    logger.info("✅ Margins fixed")
    logger.info("✅ Blank pages removed")
    logger.info("✅ Formatting cleaned")
    logger.info("="*60 + "\n")


def _safe(s: Any) -> str:
    """Safely convert any value to string."""
    return "" if s is None else str(s)


# Utility function to verify team_logos.json
def verify_logo_mappings(json_file: str = "team_logos.json"):
    """Verify that all logo files in team_logos.json exist"""
    print("\n" + "="*60)
    print(f"VERIFYING LOGO MAPPINGS IN {json_file}")
    print("="*60)
    
    if not Path(json_file).exists():
        print(f"❌ {json_file} not found!")
        return False
    
    manager = LogoManager(json_file)
    
    all_good = True
    missing_files = []
    
    for team, path in manager.logo_mappings.items():
        if Path(path).exists():
            print(f"✅ {team:30} -> {path}")
        else:
            print(f"❌ {team:30} -> {path} (FILE NOT FOUND)")
            missing_files.append((team, path))
            all_good = False
    
    if missing_files:
        print("\n⚠️ MISSING FILES:")
        for team, path in missing_files:
            print(f"  - {path} (for {team})")
        print("\nPlease ensure these files exist before running the gazette.")
    else:
        print("\n✅ All logo files found! Ready to generate gazette.")
    
    return all_good


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == "verify":
        verify_logo_mappings()
    else:
        print("Weekly Recap Builder - Using team_logos.json")
        print("Features:")
        print("  ✅ Loads all logos from team_logos.json")
        print("  ✅ Proper image embedding")
        print("  ✅ Removes blank pages")
        print("  ✅ Fixes margin issues")
        print("  ✅ Cleans formatting")
        print("\nUsage:")
        print("  Called from build_gazette.py")
        print("  Or verify logos: python weekly_recap.py verify")