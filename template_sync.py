#!/usr/bin/env python3
"""
Template Synchronization Checker for Gridiron Gazette
Ensures your template and code are properly aligned
"""

import os
import zipfile
from pathlib import Path
from docx import Document
import xml.etree.ElementTree as ET


def analyze_template(template_path: str = "recap_template.docx"):
    """
    Analyze the template to understand what placeholders it expects
    """
    print("\n" + "="*60)
    print("TEMPLATE ANALYSIS")
    print("="*60)
    
    if not Path(template_path).exists():
        print(f"❌ Template not found: {template_path}")
        return
    
    # Extract and analyze the template
    temp_dir = Path("temp_template_analysis")
    temp_dir.mkdir(exist_ok=True)
    
    try:
        # Unzip the docx
        with zipfile.ZipFile(template_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Read document.xml
        doc_xml_path = temp_dir / "word" / "document.xml"
        if doc_xml_path.exists():
            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Find all placeholders
            import re
            
            # Look for {{ }} style placeholders
            double_brace = re.findall(r'\{\{([^}]+)\}\}', content)
            
            # Look for {{% %}} style (jinja2/docxtpl)
            jinja_style = re.findall(r'\{%([^%]+)%\}', content)
            
            # Look for image placeholders
            image_refs = re.findall(r'MATCHUP\d+_[HOME|AWAY]*_LOGO', content)
            
            print("\n📝 Template Placeholders Found:")
            print("-" * 40)
            
            if double_brace:
                print("\nDouble Brace {{ }} Placeholders:")
                for placeholder in sorted(set(double_brace)):
                    print(f"  • {placeholder.strip()}")
            
            if jinja_style:
                print("\nJinja2 Style {{% %}} Tags:")
                for tag in sorted(set(jinja_style)):
                    print(f"  • {tag.strip()}")
            
            if image_refs:
                print("\nImage References:")
                for img_ref in sorted(set(image_refs)):
                    print(f"  • {img_ref}")
            
            # Check for media folder
            media_dir = temp_dir / "word" / "media"
            if media_dir.exists():
                images = list(media_dir.glob("*"))
                print(f"\n🖼️ Media files in template: {len(images)}")
                for img in images[:5]:  # Show first 5
                    print(f"  • {img.name}")
            
            # Check relationships for image handling
            rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
            if rels_path.exists():
                tree = ET.parse(rels_path)
                root = tree.getroot()
                image_rels = root.findall(".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/image']")
                print(f"\n🔗 Image relationships: {len(image_rels)}")
        
    finally:
        # Clean up temp directory
        import shutil
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
    
    print("\n" + "="*60)
    print("REQUIRED CONTEXT KEYS")
    print("="*60)
    
    # List the keys your code should provide
    required_keys = [
        "WEEK_NUMBER", "YEAR", "LEAGUE_NAME",
        "LEAGUE_LOGO", "SPONSOR_LOGO",
    ]
    
    for i in range(1, 8):
        required_keys.extend([
            f"MATCHUP{i}_HOME",
            f"MATCHUP{i}_AWAY", 
            f"MATCHUP{i}_HS",
            f"MATCHUP{i}_AS",
            f"MATCHUP{i}_HOME_LOGO",
            f"MATCHUP{i}_AWAY_LOGO",
            f"MATCHUP{i}_BLURB",
            f"MATCHUP{i}_TOP_HOME",
            f"MATCHUP{i}_TOP_AWAY",
            f"MATCHUP{i}_BUST",
            f"MATCHUP{i}_KEYPLAY",
            f"MATCHUP{i}_DEF"
        ])
    
    required_keys.extend([
        "AWARD_CUPCAKE_TEAM", "AWARD_CUPCAKE_NOTE",
        "AWARD_KITTY_TEAM", "AWARD_KITTY_NOTE",
        "AWARD_TOP_TEAM", "AWARD_TOP_NOTE"
    ])
    
    print("Your code should provide these keys:")
    for key in required_keys[:20]:  # Show first 20
        print(f"  • {key}")
    print(f"  ... and {len(required_keys) - 20} more")


def check_output_format(output_path: str):
    """
    Check the generated output for common issues
    """
    print("\n" + "="*60)
    print("OUTPUT ANALYSIS")
    print("="*60)
    
    if not Path(output_path).exists():
        print(f"❌ Output file not found: {output_path}")
        return
    
    doc = Document(output_path)
    
    # Check for markdown artifacts
    markdown_found = False
    for para in doc.paragraphs:
        if '![' in para.text or '**' in para.text or '##' in para.text:
            if not markdown_found:
                print("⚠️ Markdown syntax found in output:")
                markdown_found = True
            print(f"  • {para.text[:50]}...")
    
    if not markdown_found:
        print("✅ No markdown artifacts in output")
    
    # Check for images
    image_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if hasattr(run, '_element'):
                if run._element.xpath('.//w:drawing'):
                    image_count += 1
    
    print(f"\n🖼️ Embedded images found: {image_count}")
    
    # Check sections and margins
    print(f"\n📏 Document sections: {len(doc.sections)}")
    for i, section in enumerate(doc.sections):
        print(f"  Section {i+1}:")
        print(f"    Top margin: {section.top_margin / 914400:.2f} inches")
        print(f"    Bottom margin: {section.bottom_margin / 914400:.2f} inches")
        print(f"    Left margin: {section.left_margin / 914400:.2f} inches")
        print(f"    Right margin: {section.right_margin / 914400:.2f} inches")
    
    # Check for blank pages (excessive paragraphs)
    empty_para_count = 0
    consecutive_empty = 0
    max_consecutive = 0
    
    for para in doc.paragraphs:
        if not para.text.strip():
            empty_para_count += 1
            consecutive_empty += 1
            max_consecutive = max(max_consecutive, consecutive_empty)
        else:
            consecutive_empty = 0
    
    print(f"\n📄 Paragraph analysis:")
    print(f"  Total paragraphs: {len(doc.paragraphs)}")
    print(f"  Empty paragraphs: {empty_para_count}")
    print(f"  Max consecutive empty: {max_consecutive}")
    
    if max_consecutive > 3:
        print("  ⚠️ Excessive empty paragraphs may cause blank pages")


def generate_fix_recommendations(template_path: str, output_path: str):
    """
    Generate specific recommendations based on analysis
    """
    print("\n" + "="*60)
    print("RECOMMENDED FIXES")
    print("="*60)
    
    print("""
1. IMAGE EMBEDDING FIX:
   The template expects InlineImage objects, not paths.
   Ensure weekly_recap.py uses:
   
   context[f"MATCHUP{i}_HOME_LOGO"] = InlineImage(
       doc_template, 
       str(logo_path), 
       width=Mm(15)
   )

2. MARKDOWN REMOVAL:
   All text must be cleaned before rendering:
   
   from storymaker import clean_markdown_for_docx
   context = clean_all_markdown_in_dict(context)

3. MARGIN CONSISTENCY:
   Post-processing must set:
   - Top/Bottom: 1.0 inches (for header/footer space)
   - Left/Right: 0.75 inches
   - Header distance: 0.5 inches
   - Footer distance: 0.5 inches

4. TEMPLATE TYPE:
   Your template appears to use docxtpl syntax.
   Ensure you're using DocxTemplate, not Document:
   
   from docxtpl import DocxTemplate
   doc = DocxTemplate(template_path)
   doc.render(context)  # Not doc.save()

5. BLANK PAGE PREVENTION:
   Remove consecutive empty paragraphs in post-processing.
   Keep maximum 1 empty paragraph between sections.
""")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        if sys.argv[1] == "check":
            # Check both template and output
            template = "recap_template.docx"
            output = sys.argv[2] if len(sys.argv) > 2 else "recaps/Gazette_2025_W03.docx"
            
            analyze_template(template)
            check_output_format(output)
            generate_fix_recommendations(template, output)
        else:
            output_file = sys.argv[1]
            check_output_format(output_file)
    else:
        print("Template and Output Synchronization Checker")
        print("\nUsage:")
        print("  python template_sync.py check [output.docx]  - Full analysis")
        print("  python template_sync.py output.docx          - Check output only")
        print("\nThis tool helps identify mismatches between your template and code")